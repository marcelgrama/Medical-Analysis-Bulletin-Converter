import os
import sys
import re
import pdfplumber
from docx import Document
import pandas as pd

def remove_sensitive_data(text):
    if not isinstance(text, str):
        return text

    # Remove names using a simple pattern (capitalized words)
    text = re.sub(r'\b[A-Z][a-z]+\b', '', text)

    # Remove email addresses
    text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '', text)

    return text

def clean_pdf(pdf_path):
    cleaned_text = ""

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            cleaned_text += remove_sensitive_data(text) + "\n"

    return cleaned_text

def extract_data(text):
    data = {}
    lines = text.split("\n")

    for line in lines:
        match = re.match(r'^(.+?)\s*=\s*([\d.]+)\s*(.+?)\s*\[([\d.-]+)\s*-\s*([\d.-]+)\]', line)
        if match:
            data[match.group(1)] = {
                'value': float(match.group(2)),
                'unit': match.group(3),
                'range_low': float(match.group(4)),
                'range_high': float(match.group(5))
            }
    return data

def create_docx(cleaned_data, output_path):
    doc = Document()

    for key, value in cleaned_data.items():
        para = doc.add_paragraph()
        para.add_run(f"{key} = {value['value']} {value['unit']} [ {value['range_low']} - {value['range_high']} ]")

    doc.save(output_path)

def create_csv(cleaned_data, output_path):
    df = pd.DataFrame(cleaned_data).T.reset_index()
    df.columns = ['Parameter', 'Value', 'Unit', 'Range Low', 'Range High']
    df.to_csv(output_path, index=False)

if __name__ == "__main__":
    pdf_path = sys.argv[1]

    with pdfplumber.open(pdf_path) as pdf:
        text = "\n".join(page.extract_text() for page in pdf.pages)

    cleaned_text = remove_sensitive_data(text)
    cleaned_data = extract_data(cleaned_text)

    docx_path = os.path.splitext(pdf_path)[0] + "_cleaned.docx"
    create_docx(cleaned_data, docx_path)

    csv_path = os.path.splitext(pdf_path)[0] + "_cleaned.csv"
    create_csv(cleaned_data, csv_path)
